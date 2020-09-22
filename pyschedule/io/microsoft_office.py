# -*- coding: utf-8 -*-
"""
@date Thu Sep  3 10:30:27 2020

@brief IO for reading/writing from microsoft office formats

@author: aweis
"""

import copy
import numpy as np
import datetime

try:
    import openpyxl
except ModuleNotFoundError:
    print("Cannot import OpenPyxl (no Excel capability). Try `pip install openpyxl`.")
    
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.style import WD_STYLE_TYPE
except ModuleNotFoundError:
    print("Cannot import docx (no Word capability). Try `pip install python-docx`.")
    
from pyschedule.Tasks import unpack_children_with_levels
from pyschedule.Schedule import Schedule,get_range_ymd,get_datetime_range

#%% Write schedule to excel

def write_excel(schedule:Schedule,fpath:str,**kwargs):
    '''
    @brief create a gantt chart from an excel spreadsheet
    @param[in] schedule - schedule to write out
    @param[in] fpath - path to file to write spreadsheet to
    '''
    wb = openpyxl.Workbook()
    ws = wb.active
    
    # Lets first load everything so we know the sizes of tasks, times, etc
    ## Load our tasks
    tasks,levels = unpack_children_with_levels(schedule,-1)
    task_levels = max(levels)-min(levels)+1
    ## Now get date/time range 
    years,months,days = get_range_ymd(schedule.start, schedule.end)
    date_range = get_datetime_range(schedule.start,schedule.end)
    ym = [date.strftime('%Y-%m') for date in date_range]
    ymu = np.unique(ym)
    time_cols = np.size(months)
    
    # Set up the formatting
    ## Header Formatting
    #head_fmt = wb.add_format({'bold':True,'align':'center'})
    ## Task formatting
    #task_col_fmt = wb.add_format({'text_wrap':True})
    ## Date Header Formatting
    #date_head_fmt = head_fmt
    
    # Now lets set the sizes of everything
    ## Our header
    head_row_start = 1; 
    head_col_start = 1
    head_rows = 2; 
    head_cols = task_levels 
    head_row_end = head_row_start+head_rows-1
    head_col_end = head_col_start+head_cols-1
    ## Our tasks
    task_row_start = head_row_start+head_rows;
    task_col_start = 1
    task_rows = len(tasks)
    task_cols = task_levels 
    task_row_end = task_row_start+task_rows-1
    task_col_end = task_col_start+task_cols-1
    ## Date header
    date_head_row_start = 1
    date_head_col_start = head_col_end+1
    date_head_rows = head_rows 
    date_head_cols = sum([len(month) for month in months])
    ## Date range
    date_range_row_start = date_head_row_start+date_head_rows
    date_range_col_start = date_head_col_start 
    date_range_rows = task_rows
    date_range_cols = date_head_cols
    ## Totals
    total_row_start = 1
    total_col_start = 1
    total_rows = head_rows+task_rows
    total_cols = date_head_col_start+date_head_cols
    total_row_end = total_row_start+total_rows-1
    total_col_end = total_col_start+total_cols-1
    
    # now write out values
    ## Write our header
    ws.cell(row=head_row_start,column=head_col_start,value="Tasks")
    ## Write our date header
    date_head_col = date_head_col_start; date_head_row = date_head_row_start
    for yi,y in enumerate(years):
        ws.cell(row=date_head_row,column=date_head_col,value=y)
        ws.merge_cells(start_row=date_head_row,start_column=date_head_col,
                       end_row=date_head_row,end_column=date_head_col+len(months[yi])-1)
        month_col = date_head_col
        for month in months[yi]:
            ws.cell(row=date_head_row+1,column=month_col,value=month)
            month_col += 1
        date_head_col += len(months[yi])
    ## Write out our tasks
    task_row=task_row_start; task_col=task_col_start
    for t,l in zip(tasks,levels):
        mycell = ws.cell(row=task_row,column=task_col+l,value=t.nickname)
        mycell.alignment = openpyxl.styles.Alignment(wrap_text=True)
        #print(' '*2*l+t['name'])
        task_row += 1
    for c in range(task_col_start,task_col_end+1): 
        ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width = 20
        
    # Now shade the columns for the tasks
    for task_row,t in enumerate(tasks):
        td_range = get_datetime_range(t.start,t.end)
        tdu_range = np.unique([date.strftime('%Y-%m') for date in td_range])
        trows = [date_range_row_start+task_row]
        tcols = [date_range_col_start+np.where(ymu==tdu)[0][0] for tdu in tdu_range]
        for r in trows:
            for c in tcols:
                cell = ws.cell(row=r,column=c)
                shading = openpyxl.styles.PatternFill("solid", fgColor="88888888")
                cell.fill = shading
                
    
    # Draw our borders but dont change other formatting
    thick_border = openpyxl.styles.Side(border_style="thick", color="000000")
    ## Bottom borders
    bb_rows = [head_row_end,total_row_end]+list(task_row_start-1+np.where(np.array(levels)==0)[0])
    bb_cols = range(total_col_start,total_col_end)
    for r in bb_rows:
        for c in bb_cols:
            cell = ws.cell(row=r,column=c)
            myborder = copy.copy(cell.border)
            myborder.bottom = thick_border
            cell.border = myborder
    # Right borders
    br_cols = [task_col_end]+list(date_head_col_start-1+np.cumsum([len(m) for m in months]))
    br_rows = range(total_row_start,total_row_end+1)
    for r in br_rows:
        for c in br_cols:
            cell = ws.cell(row=r,column=c)
            myborder = copy.copy(cell.border)
            myborder.right = thick_border
            cell.border = myborder
    # Draw borders for 'today'
    today_str = datetime.datetime.today().strftime('%Y-%m')
    today_col = np.where(np.array(ymu)==today_str)[0][0]+date_range_col_start
    today_border = openpyxl.styles.Side(border_style="thick", color="00FF00")
    for r in range(total_row_start,total_row_end+1):
        cell = ws.cell(row=r,column=today_col)
        myborder = copy.copy(cell.border)
        myborder.right = today_border
        myborder.left = today_border
        cell.border = myborder
    
    # finally save it out
    wb.save(filename = fpath)
    
    return wb
    

#%% Write schedule out to word

def write_word(schedule:Schedule,fpath):
    '''
    @brief Write out a schedule to a word file
    @param[in] schedule - schedule to write out
    @param[in] fpath - path to file to write spreadsheet to
    '''
    
    # Initialize the document and write title info
    doc = Document()
    title = doc.add_paragraph(schedule['name'],style='Title')
    
    # Add some top level info
    for k in ['description']:
        val = schedule.get(k,None)
        if val not in schedule.UNDEFINED_VALS:
            par = doc.add_paragraph('{}'.format(k.capitalize()),style='Heading 1')
            doc.add_paragraph(str(val),style='Normal')        
    
    # Load in some task info
    tasks,levels = unpack_children_with_levels(schedule,-1)
    
    # add our styles
    for level in np.unique(levels):
        ## Add subparagraph style
        task_heading_style = doc.styles.add_style('task_par_{}'.format(level),WD_STYLE_TYPE.PARAGRAPH)
        task_heading_style.base_style = doc.styles['Normal']
        task_heading_style.paragraph_format.left_indent = Inches(0.25*level)
        ## Add a header style
        task_heading_style = doc.styles.add_style('task_heading_{}'.format(level),WD_STYLE_TYPE.PARAGRAPH)
        task_heading_style.base_style = doc.styles['Heading {}'.format(level+2)] #+2 to start at heading 2
        task_heading_style.paragraph_format.left_indent = Inches(0.25*level)
        task_heading_style.next_paragraph_style = doc.styles['task_par_{}'.format(level)]
    
    # now iterate and write out tasks
    doc.add_paragraph('Task List',style='Heading 1')
    for task,level in zip(tasks,levels):
        level = level
        ## write the header
        task_head = doc.add_paragraph(task['name'],style='task_heading_{}'.format(level))
        ## Write some static properties
        for k in ['description']:
            val = task.get(k,None)
            if val not in task.UNDEFINED_VALS:
                par = doc.add_paragraph(style='task_par_{}'.format(level))
                par.add_run('{} :'.format(k.capitalize())).bold=True 
                par.add_run(str(val))
        ## Write dynamic properties
        for k in ['progress']:
            val = getattr(task,k)
            if val not in task.UNDEFINED_VALS:
                par = doc.add_paragraph(style='task_par_{}'.format(level))
                par.add_run('{} :'.format(k.capitalize())).bold=True 
                par.add_run(str(val))
        ## Add timeframe
        ptime = doc.add_paragraph(style='task_par_{}'.format(level))
        ptime.add_run('Expected Timeframe: ').bold=True
        ptime.add_run('{} - {} ({} days)'.format(task.start.strftime('%Y-%m-%d'),
                                            task.end.strftime('%Y-%m-%d'),
                                            task.duration.days))
        
    doc.save(fpath)
    
    return doc
    
    
#%% Some testing
if __name__=='__main__':
    
    mydoc = Document()
    
    for level in [0,1,2,3,4]:
        
        mydoc.add_paragraph('Head{}'.format(level),style='Heading {}'.format(level+1))
        mydoc.add_paragraph('test paragraph level {}'.format(level))
    






